from flask import Flask, render_template, request, redirect, url_for, session, send_file
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from werkzeug.utils import secure_filename
import os

app = Flask(__name__)
app.secret_key = 'your_secret_key'

# Sample user data (replace with your own database or user management system)
users = {
    'admin': 'admin123'
}

# Restricted usernames
restricted_usernames = ['admin', 'root', 'superuser']

# Define the allowed file extensions
ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif'}

# Set the upload folder
app.config['UPLOAD_FOLDER'] = 'uploads'

# Function to check if a filename has an allowed extension
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        # Check username and password
        if username in users and users[username] == password:
            session['username'] = username  # Store username in session
            return redirect(url_for('dashboard'))  # Redirect to dashboard after successful login
        else:
            return render_template('login.html', error='Invalid username or password')
    return render_template('login.html')

@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST': 
        username = request.form['username']
        password = request.form['password']
        
        # Check if username or password are empty
        if not username or not password:
            return render_template('register.html', error='Username and password are required')
        
        # Check if username already exists
        if username in users:
            return render_template('register.html', error='Username already exists')
        
        # Check if the username is restricted
        if username in restricted_usernames:
            return render_template('register.html', error='Username is restricted, please choose another one')
        
        # Check password strength
        if len(password) < 8:
            return render_template('register.html', error='Password must be at least 8 characters long')
        
        # Check if username contains restricted characters
        if not username.isalnum():
            return render_template('register.html', error='Username must contain only letters and numbers')
        
        # Check if password contains restricted characters
        if not password.isalnum():
            return render_template('register.html', error='Password must contain only letters and numbers')
        
        # Add new user
        users[username] = password
        session['username'] = username  # Store username in session
        return redirect(url_for('dashboard'))  # Redirect to dashboard after registration
    
    return render_template('register.html')

@app.route('/dashboard')
def dashboard():
    return render_template('dashboard.html')

@app.route('/logout')
def logout():
    session.pop('username', None)  # Remove username from session
    return redirect(url_for('index'))  # Redirect to the index page after logout

@app.route('/document_request', methods=['GET', 'POST'])
def document_request():
    if request.method == 'POST':
        document_type = request.form['document_type']
        if document_type == 'barangay_clearance':
            return redirect(url_for('barangay_clearance_request'))
        elif document_type == 'residence_certification':
            return redirect(url_for('residence_certification_request'))
        elif document_type == 'indigency':
            return redirect(url_for('indigency_request'))
        else:
            # Handle invalid document types
            return "Invalid document type"
    return render_template('document_request.html')
    
def validate_barangay_clearance(form_data):
    # Add your validation logic here
    # For simplicity, let's just check if all required fields are filled
    required_fields = ['first_name', 'middle_name', 'last_name', 'address', 'purpose']
    for field in required_fields:
        if field not in form_data or not form_data[field]:
            return False
    
    # Check if photo is uploaded
    if 'photo' not in request.files or request.files['photo'].filename == '':
        return False
    
    return True

@app.route('/barangay_clearance_request', methods=['GET', 'POST'])
def barangay_clearance_request():
    if request.method == 'POST':
        # Perform validation
        is_valid = validate_barangay_clearance(request.form)
        if is_valid:
            # Store form data in session
            session['form_data'] = request.form
            # Handle photo upload
            photo = request.files['photo']
            if photo.filename != '':
                if allowed_file(photo.filename):
                    # Save the uploaded photo temporarily
                    photo_path = os.path.join(app.config['UPLOAD_FOLDER'], secure_filename(photo.filename))
                    photo.save(photo_path)
                    session['photo_path'] = photo_path
                else:
                    return render_template('barangay_clearance_form.html', error="Invalid file format. Allowed formats: png, jpg, jpeg, gif", form_data=request.form)
            else:
                session['photo_path'] = None
            # Redirect to document generation
            return redirect(url_for('document_generated'))
        else:
            # Display error message and prompt for corrections
            return render_template('barangay_clearance_form.html', error="Invalid information. Please correct the fields.", form_data=request.form)
    return render_template('barangay_clearance_form.html')

@app.route('/indigency_request', methods=['GET', 'POST'])
def indigency_request():
    if request.method == 'POST':
        # Perform validation
        is_valid = validate_indigency(request.form)
        if is_valid:
            # Store form data in session
            session['form_data'] = request.form
            # Redirect to document generation or any other desired endpoint
            return redirect(url_for('document_generated'))
        else:
            # Display error message and prompt for corrections
            return render_template('indigency.html', error="Invalid information. Please correct the fields.", form_data=request.form)
    return render_template('indigency.html')

def validate_indigency(form_data):
    # Add your validation logic here
    # For simplicity, let's just check if all fields are filled
    for value in form_data.values():
        if not value:
            return False
    return True

from PIL import Image


def generate_document(form_data, logo_left_path=None, logo_right_path=None, photo_path=None):
    # Create a new Word document
    doc = Document()

    # Add the title with logos beside it
    title_with_logos = doc.add_paragraph()
    title_with_logos.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add left logo
    if logo_left_path:
        run_left = title_with_logos.add_run()
        left_logo_shape = run_left.add_picture(logo_left_path, width=Inches(0.75))
        left_logo_shape.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add title
    title_run = title_with_logos.add_run("Republic Act of the Philippines\nBarangay Bagumbayan\nTanauan City, Batangas\n\nBARANGAY CLEARANCE\n\n")
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_run.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add right logo
    if logo_right_path:
        run_right = title_with_logos.add_run()
        right_logo_shape = run_right.add_picture(logo_right_path, width=Inches(0.75))
        right_logo_shape.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add a line break
    doc.add_paragraph()
        
    # Add thick line below the title
    thick_line = doc.add_paragraph("______________________________________________________________________________________________________________")

    # Add the "Barangay Clearance" title
    barangay_clearance_title = doc.add_paragraph()
    barangay_clearance_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    barangay_clearance_run = barangay_clearance_title.add_run("BARANGAY CLEARANCE")
    barangay_clearance_run.bold = True
    barangay_clearance_run.font.size = Pt(14)

    # Add a line break
    doc.add_paragraph()

    # Add recipient's name and details
    recipient_details = doc.add_paragraph()
    recipient_details.add_run("TO WHOM IT MAY CONCERN:\n\nThis is to certify that, based on the records of this Barangay, the person whose name and personal circumstances are indicated below has not been accused nor has a pending case with Barangay Bagumbayan involving moral turpitude or any act contrary to existing law:\n\nFirst Name: ").bold = True
    recipient_details.add_run(f"{form_data['first_name']}").bold = False
    recipient_details.add_run("\nMiddle Name: ").bold = True
    recipient_details.add_run(f"{form_data['middle_name']}").bold = False
    recipient_details.add_run("\nLast Name: ").bold = True
    recipient_details.add_run(f"{form_data['last_name']}").bold = False
    recipient_details.add_run("\nAddress: ").bold = True
    recipient_details.add_run(f"{form_data['address']}").bold = False
    recipient_details.add_run("\nPurpose: ").bold = True
    recipient_details.add_run(f"{form_data['purpose']}").bold = False

    # Add photo if path is provided
    if photo_path:
        try:
            # Open and resize the image
            img = Image.open(photo_path)
            img.thumbnail((200, 200))  # Resize the image
            img_path = 'temp_photo.jpg'
            img.save(img_path)  # Save the resized image temporarily
            
            # Add the photo to the document
            doc.add_picture(img_path, width=Inches(1))  # Adjust width as needed
            
            # Remove the temporary image file
            os.remove(img_path)
        except Exception as e:
            print(f"Error adding photo to document: {e}")

    # Add signature and date
    signature_date = doc.add_paragraph("\n\nThis certification is issued upon the request of the above-named person to support whatever legal purposes it may serve best.\n\nGiven this ____ day of __________ 20__ at Barangay Bagumbayan, Tanauan City, Batangas, Philippines.\n\nHON. **********\n\nPunong Barangay")
    signature_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    output_dir = 'output_documents'
    os.makedirs(output_dir, exist_ok=True)  # Create the directory if it doesn't exist
    output_path = os.path.join(output_dir, "certificate.docx")
    doc.save(output_path)

    print("Certificate generated successfully!")

    return output_path

@app.route('/document_generated')
def document_generated():
    # Generate document and get the file path
    file_path = generate_document(session.get('form_data', {}), session.get('photo_path'))
    
    # Display success message
    success_message = "Document Successfully Generated"
    
    # Provide a link to download the generated document
    return render_template('document_generated.html', success_message=success_message, file_path=file_path)

from flask import send_file
from io import BytesIO

@app.route('/download_file')
def download_file():
    # Retrieve form data from session
    form_data = session.get('form_data', {})
    
    # Check if form data is empty
    if not form_data:
        return "Form data not found"
    
    # Placeholder paths for logo and photo
    logo_left_path = 'static/logo.jpg'
    logo_right_path = 'static/logo.jpg'
    photo_path = session.get('photo_path')  # Get the photo path from session
    
    # Generate the document content
    document_content = generate_document(form_data, logo_left_path, logo_right_path,  photo_path)  # Pass form_data dictionary here
    
    # Convert the document content to bytes
    document_bytes = BytesIO(document_content.encode())

    # Send the document for download
    return send_file(document_bytes, as_attachment=True, download_name='barangay_clearance.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

@app.route('/user_feedback')
def user_feedback():
    # Render the user_feedback.html template
    return render_template('user_feedback.html')

@app.route('/final_output')
def final_output():
    # Render the final_output.html template
    return render_template('final_output.html')

if __name__ == '__main__':
    app.run(debug=True)